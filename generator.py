import os
import re
import glob
import shutil
from datetime import datetime
from difflib import SequenceMatcher
from docx import Document
from docx2pdf import convert
from pypdf import PdfReader, PdfWriter
from dotenv import load_dotenv

import config
from classifier import classify_job as _classify_job
from classifier import load_keywords as _classifier_load_keywords
from llm import call_llm
from resume_context import (
    build_resume_context_query,
    get_resume_context_paths as _get_resume_context_paths,
    load_resume_extended as _load_resume_extended,
    load_resume_source as _load_resume_source,
    resume_context_status as _resume_context_status_lines,
    select_resume_extended_context,
)


# =============================================================================
# Config Defaults
# =============================================================================

try:
    from config import BUNDLE_APPENDIX
except ImportError:
    BUNDLE_APPENDIX = []
try:
    from config import BUNDLE_NAME
except ImportError:
    BUNDLE_NAME = "Cover Letter Bundle"
try:
    from config import RESUME_SOURCE
except ImportError:
    RESUME_SOURCE = "resume.source.md"
try:
    from config import RESUME_SOURCE_FILENAME
except ImportError:
    RESUME_SOURCE_FILENAME = "resume.source.md"
try:
    from config import RESUME_EXTENDED_SOURCE
except ImportError:
    RESUME_EXTENDED_SOURCE = "resume.extended.md"
try:
    from config import RESUME_EXTENDED_FILENAME
except ImportError:
    RESUME_EXTENDED_FILENAME = "resume.extended.md"


# =============================================================================
# Environment And Project Paths
# =============================================================================

load_dotenv()


# =============================================================================
# Keyword Loading
# =============================================================================

def _load_keywords():
    return _classifier_load_keywords()


# =============================================================================
# Config Accessors
# =============================================================================

def get_resume_page_limit():
    try:
        limit = int(getattr(config, "RESUME_PAGE_LIMIT", 1))
    except (TypeError, ValueError):
        limit = 1
    return max(1, limit)


def get_cover_letter_page_limit():
    try:
        limit = int(getattr(config, "COVER_LETTER_PAGE_LIMIT", 1))
    except (TypeError, ValueError):
        limit = 1
    return max(1, limit)


def get_page_fit_max_attempts():
    try:
        attempts = int(getattr(config, "PAGE_FIT_MAX_ATTEMPTS", 2))
    except (TypeError, ValueError):
        attempts = 2
    return max(0, attempts)


def get_page_fit_max_lines_per_attempt():
    try:
        lines = int(getattr(config, "PAGE_FIT_MAX_LINES_PER_ATTEMPT", 4))
    except (TypeError, ValueError):
        lines = 4
    return max(1, lines)


def get_page_fit_min_line_retain_ratio():
    try:
        ratio = float(getattr(config, "PAGE_FIT_MIN_LINE_RETAIN_RATIO", 0.88))
    except (TypeError, ValueError):
        ratio = 0.88
    return min(1.0, max(0.5, ratio))


def get_resume_tailoring_aggression():
    mode = str(getattr(config, "RESUME_TAILORING_AGGRESSION", "balanced")).strip().lower()
    if mode not in {"conservative", "balanced", "aggressive"}:
        return "balanced"
    return mode


def get_cli_verbosity():
    mode = str(getattr(config, "CLI_VERBOSITY", "normal")).strip().lower()
    if mode not in {"quiet", "normal", "debug"}:
        return "normal"
    return mode


def _resume_aggression_guidance(mode):
    if mode == "conservative":
        return (
            "conservative - edit only direct, obvious matches. Prefer SKIP when "
            "the sentence-to-keyword fit is merely adjacent."
        )
    if mode == "aggressive":
        return (
            "aggressive - permit stronger adjacent phrasing when the factual "
            "source or extended source supports it, while still avoiding invented "
            "ownership, tools, credentials, or outcomes."
        )
    return (
        "balanced - edit direct matches and coherent adjacent matches. Use SKIP "
        "for weak or forced sentence-to-keyword fits."
    )


# =============================================================================
# Resume Context Files
# =============================================================================

def _resolve_optional_project_path(path):
    from resume_context import resolve_optional_project_path
    return resolve_optional_project_path(path)


def _resolve_resume_markdown_path(category_dir=None, local_filename=None, fallback_path=None):
    from resume_context import resolve_resume_markdown_path
    return resolve_resume_markdown_path(category_dir, local_filename, fallback_path)


def _load_resume_markdown(category_dir=None, local_filename=None, fallback_path=None):
    from resume_context import load_resume_markdown
    return load_resume_markdown(category_dir, local_filename, fallback_path)


def get_resume_context_paths(
    category_dir=None,
    source_filename=None,
    extended_filename=None,
    project_source=None,
    project_extended=None,
):
    return _get_resume_context_paths(
        category_dir=category_dir,
        source_filename=source_filename if source_filename is not None else RESUME_SOURCE_FILENAME,
        extended_filename=extended_filename if extended_filename is not None else RESUME_EXTENDED_FILENAME,
        project_source=project_source if project_source is not None else RESUME_SOURCE,
        project_extended=project_extended if project_extended is not None else RESUME_EXTENDED_SOURCE,
    )


def _resume_context_status(context_paths, extended_selection=None):
    return _resume_context_status_lines(context_paths, extended_selection)


def load_resume_source(path=None, category_dir=None, project_source=None, source_filename=None):
    return _load_resume_source(
        path=path,
        category_dir=category_dir,
        project_source=project_source,
        source_filename=source_filename if source_filename is not None else RESUME_SOURCE_FILENAME,
    )


def load_resume_extended(path=None, category_dir=None, project_extended=None, extended_filename=None):
    return _load_resume_extended(
        path=path,
        category_dir=category_dir,
        project_extended=project_extended,
        extended_filename=extended_filename if extended_filename is not None else RESUME_EXTENDED_FILENAME,
    )


# =============================================================================
# Text Patterns
# =============================================================================

_MONTHS = (
    r"January|February|March|April|May|June|July|August|September|October|November|December"
    r"|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec"
)
_ORD = r"(?:st|nd|rd|th)?"   # optional ordinal suffix
DATE_PATTERN = re.compile(
    rf"\d{{1,2}}{_ORD}\s+(?:{_MONTHS})\s+20\d{{2}}"   # DD[th] Month YYYY  (e.g. 13th February 2026)
    rf"|(?:{_MONTHS})\s+\d{{1,2}}{_ORD},?\s+20\d{{2}}" # Month DD[th], YYYY (e.g. April 15th, 2026)
    rf"|\d{{1,2}}/\d{{1,2}}/20\d{{2}}"                 # DD/MM/YYYY         (e.g. 15/04/2026)
    rf"|20\d{{2}}-\d{{2}}-\d{{2}}",                    # YYYY-MM-DD         (e.g. 2026-04-15)
    re.IGNORECASE,
)

RESUME_MARKER_PATTERN = re.compile(r"^\s*\[(.+)\]\s*$")
RESUME_CONTROL_LABEL_PATTERN = re.compile(r"^[A-Z_]+(?::|\Z)")


# Australian states/territories and common location suffixes to strip from job titles
_LOCATION_TOKENS = re.compile(
    r"[\s,|\-]+\b("
    r"NSW|VIC|QLD|SA|WA|TAS|NT|ACT"          # AU states/territories
    r"|Australia|Remote|Hybrid|On[- ]?site"   # other common tags
    r")\b[\s,|\-]*$",
    re.IGNORECASE,
)


# =============================================================================
# DOCX/Text Helpers
# =============================================================================

def _iter_all_paragraphs(doc):
    """Yield every paragraph in the document, including those inside table cells."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def clean_job_title(title: str) -> str:
    """Strip trailing location/unnecessary tokens from a job title.

    Example: 'Investment Analyst - Investments & Treasury NSW' -> 'Investment Analyst - Investments & Treasury'
    """
    return _normalize_title_text(_LOCATION_TOKENS.sub("", title)).strip()


_FORBIDDEN_DASH_PATTERN = re.compile(r"\s*(?:[\u2013\u2014\u2212]|--+)\s*")


def _replace_forbidden_dashes(text, replacement=" - "):
    cleaned = _FORBIDDEN_DASH_PATTERN.sub(replacement, str(text or ""))
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    cleaned = re.sub(r"\s+([,.;:!?])", r"\1", cleaned)
    return cleaned


def _normalize_title_text(text):
    return _replace_forbidden_dashes(text, " - ").strip()


def _normalize_generated_text(text):
    return _replace_forbidden_dashes(text, " - ")


def _set_para_text(para, text):
    """Replace all visible paragraph text without rebuilding paragraph XML."""
    _replace_text_with_diff_preserving_runs(
        para,
        0,
        para.text,
        _normalize_generated_text(text),
    )


def _replace_text_once_preserving_runs(para, old_text, new_text):
    start = para.text.find(old_text)
    if start < 0:
        return False
    _replace_text_with_diff_preserving_runs(
        para,
        start,
        old_text,
        _normalize_generated_text(new_text),
    )
    return True


def _delete_text_once_preserving_runs(para, old_text):
    current = para.text
    start = current.find(old_text)
    if start < 0:
        return False
    end = start + len(old_text)
    if start > 0 and current[start - 1].isspace():
        start -= 1
    elif end < len(current) and current[end].isspace():
        end += 1
    _replace_para_text_range(para, start, end, "")
    return True


def _replace_text_with_diff_preserving_runs(para, offset, old_text, new_text):
    matcher = SequenceMatcher(None, old_text, new_text, autojunk=False)
    for tag, i1, i2, j1, j2 in reversed(matcher.get_opcodes()):
        if tag == "equal":
            continue
        _replace_para_text_range(para, offset + i1, offset + i2, new_text[j1:j2])


def _replace_para_text_range(para, start, end, replacement):
    replacement = _normalize_generated_text(replacement)
    text_length = len(para.text)
    start = max(0, min(start, text_length))
    end = max(start, min(end, text_length))
    runs = list(para.runs)
    if not runs:
        para.add_run(replacement)
        return

    if start == end:
        pos = 0
        for run in runs:
            run_text = run.text
            run_end = pos + len(run_text)
            if pos <= start <= run_end:
                local = start - pos
                run.text = f"{run_text[:local]}{replacement}{run_text[local:]}"
                return
            pos = run_end
        runs[-1].text += replacement
        return

    pos = 0
    inserted = False
    for run in runs:
        run_text = run.text
        run_start = pos
        run_end = pos + len(run_text)
        pos = run_end
        if run_end <= start or run_start >= end:
            continue

        local_start = max(start - run_start, 0)
        local_end = min(end - run_start, len(run_text))
        before = run_text[:local_start]
        after = run_text[local_end:]
        if not inserted:
            run.text = before + replacement + (after if end <= run_end else "")
            inserted = True
        else:
            run.text = after if end <= run_end else ""


# =============================================================================
# Job Classification
# =============================================================================

def classify_job(title, description):
    return _classify_job(
        title,
        description,
        template_base=config.TEMPLATE_BASE,
        keyword_loader=_load_keywords,
    )


# =============================================================================
# Template Discovery And Output Names
# =============================================================================

def get_paths(category):
    base = os.path.join(config.TEMPLATE_BASE, category)
    if not os.path.isdir(base):
        raise FileNotFoundError(f"Template folder not found: {base}")

    resume_pdf = _find_configured_template_file(
        base,
        "RESUME_ORIGINAL_PDF_GLOB",
        "resume PDF",
    )
    resume_docx = _find_configured_template_file(
        base,
        "RESUME_EDITABLE_DOCX_GLOB",
        "editable resume DOCX",
    )
    cover_docx = _find_configured_template_file(
        base,
        "COVER_LETTER_DOCX_GLOB",
        "cover letter DOCX",
    )

    resume_docx_candidates = []
    for f in os.listdir(base):
        if f.startswith("~$"):          # skip Word temporary lock files
            continue
        f_lower = f.lower()
        if resume_pdf is None and f_lower.endswith(".pdf") and "resume" in f_lower:
            resume_pdf = os.path.join(base, f)
        elif resume_docx is None and f_lower.endswith(".docx") and "resume" in f_lower:
            resume_docx_candidates.append(os.path.join(base, f))
        elif cover_docx is None and f_lower.endswith(".docx") and "cover" in f_lower:
            cover_docx = os.path.join(base, f)

    def _resume_docx_rank(path):
        name = os.path.basename(path).lower()
        draft_penalty = sum(
            token in name
            for token in ("clean", "test", "example", "editable", "copy")
        )
        primary_name = (
            name == "resume.docx"
            or name.endswith("_resume.docx")
            or name.endswith(" resume.docx")
            or name.endswith("-resume.docx")
        )
        return (draft_penalty, 0 if primary_name else 1, name)

    resume_docx = resume_docx or (
        sorted(resume_docx_candidates, key=_resume_docx_rank)[0]
        if resume_docx_candidates else None
    )

    missing = [
        name for name, val in [
            ("Resume.pdf", resume_pdf),
            ("Cover Letter.docx", cover_docx),
        ]
        if val is None
    ]
    if missing:
        raise FileNotFoundError(f"Missing in {base}: {', '.join(missing)}")
    return resume_pdf, resume_docx, cover_docx, base


def _find_configured_template_file(base, config_name, label):
    pattern = getattr(config, config_name, None)
    if pattern is None:
        return None

    pattern = str(pattern).strip()
    if not pattern:
        return None

    search_pattern = pattern if os.path.isabs(pattern) else os.path.join(base, pattern)
    matches = sorted(
        path for path in glob.glob(search_pattern)
        if os.path.isfile(path) and not os.path.basename(path).startswith("~$")
    )
    if not matches:
        raise FileNotFoundError(
            f"Missing configured {label} matching {pattern!r} in {base}"
        )
    return matches[0]


def get_tailored_resume_pdf_path(resume_docx_dest):
    folder = os.path.dirname(resume_docx_dest)
    resume_stem = os.path.splitext(os.path.basename(resume_docx_dest))[0]
    resume_stem_clean = _clean_resume_edit_suffix(resume_stem)
    name_template = getattr(
        config,
        "RESUME_TAILORED_PDF_NAME",
        "{resume_stem_clean}.pdf",
    )

    if name_template is None:
        filename = f"{resume_stem}.pdf"
    else:
        name_template = str(name_template).strip() or "{resume_stem_clean}.pdf"
        filename = name_template.format(
            resume_stem=resume_stem,
            resume_stem_clean=resume_stem_clean,
        )
        if not filename.lower().endswith(".pdf"):
            filename = f"{filename}.pdf"

    return os.path.join(folder, os.path.basename(filename))


def get_cover_letter_pdf_path(cover_docx_dest):
    folder = os.path.dirname(cover_docx_dest)
    cover_stem = os.path.splitext(os.path.basename(cover_docx_dest))[0]
    cover_stem_clean = _clean_resume_edit_suffix(cover_stem)
    name_template = getattr(
        config,
        "COVER_LETTER_PDF_NAME",
        "{cover_stem_clean}.pdf",
    )

    if name_template is None:
        filename = f"{cover_stem}.pdf"
    else:
        name_template = str(name_template).strip() or "{cover_stem_clean}.pdf"
        filename = name_template.format(
            cover_stem=cover_stem,
            cover_stem_clean=cover_stem_clean,
        )
        if not filename.lower().endswith(".pdf"):
            filename = f"{filename}.pdf"

    return os.path.join(folder, os.path.basename(filename))


def _clean_resume_edit_suffix(stem):
    return re.sub(r"([ _-])edit(?:able)?$", "", stem, flags=re.IGNORECASE)


# =============================================================================
# Cover Letter Filling
# =============================================================================

_COVER_CONTROL_LABEL_PATTERN = re.compile(
    r"(?:^|[\s.])\[?\s*(?:OPTIONAL|DESCRIPTION)\s*:",
    flags=re.IGNORECASE,
)


def _has_cover_letter_blank(text):
    return "_" in text or "[" in text or bool(_COVER_CONTROL_LABEL_PATTERN.search(text))


def _is_optional_cover_marker(text):
    return re.search(r"\[?\s*OPTIONAL\s*:", text, flags=re.IGNORECASE) is not None


def _is_cover_letter_delete_rewrite(text):
    normalized = re.sub(r"[^A-Z]", "", text.upper())
    return normalized in {"DELETE", "REMOVE", "SKIP"}


def _split_cover_letter_sentences(text):
    """Split cover-letter text without splitting inside bracketed instructions."""
    sentences = []
    start = 0
    depth = 0
    i = 0
    while i < len(text):
        ch = text[i]
        if ch == "[":
            depth += 1
        elif ch == "]" and depth > 0:
            depth -= 1
        elif depth == 0 and ch in ".!?":
            j = i + 1
            while j < len(text) and text[j].isspace():
                j += 1
            if j < len(text) and (text[j].isupper() or text[j] == "["):
                sentence = text[start:i + 1].strip()
                if sentence:
                    sentences.append(sentence)
                start = j
                i = j
                continue
        i += 1

    tail = text[start:].strip()
    if tail:
        sentences.append(tail)
    return sentences


def _clean_cover_letter_rewrite(text):
    """Remove leaked square-bracket template instructions from an LLM rewrite."""
    if not text:
        return ""

    output = []
    depth = 0
    for ch in text:
        if ch == "[":
            depth += 1
            continue
        if ch == "]" and depth > 0:
            depth -= 1
            continue
        if depth == 0:
            output.append(ch)

    cleaned = "".join(output)
    cleaned = re.sub(r"\s*\bOPTIONAL:\s*.*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*\bDESCRIPTION:\s*.*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    cleaned = re.sub(r"\s+([,.;:!?])", r"\1", cleaned)
    return _normalize_generated_text(cleaned)


def _prepare_cover_letter_rewrite(raw_sentence, original_sentence):
    raw_sentence = (raw_sentence or "").strip()
    if _is_cover_letter_delete_rewrite(raw_sentence):
        return "", True

    if _is_optional_cover_marker(original_sentence):
        payload = _optional_rewrite_payload(raw_sentence)
        candidate_source = payload if payload is not None else raw_sentence
        candidate = _extract_optional_final_sentence(candidate_source)
        candidate = _clean_cover_letter_rewrite(candidate)
        if not candidate or _looks_like_unfilled_optional_instruction(candidate):
            return "", True
        return _ensure_sentence_punctuation(candidate), False

    return _clean_cover_letter_rewrite(raw_sentence), False


def _optional_rewrite_payload(text):
    cleaned = (text or "").strip()
    if cleaned.startswith("["):
        cleaned = cleaned[1:].strip()
    if not re.match(r"^OPTIONAL\s*:", cleaned, flags=re.IGNORECASE):
        return None

    cleaned = re.sub(r"^OPTIONAL\s*:\s*", "", cleaned, flags=re.IGNORECASE).strip()
    if cleaned.endswith("]"):
        cleaned = cleaned[:-1].strip()
    return cleaned


def _extract_optional_final_sentence(text):
    cleaned = (text or "").strip()

    quoted = re.findall(r'["“](.+?)["”]', cleaned)
    for item in quoted:
        if not _looks_like_unfilled_optional_instruction(item):
            return item.strip()

    sentence_match = re.search(
        r"\b(I\s+(?:have|bring|developed|built|used|supported|worked|managed|led|can)\b.+)",
        cleaned,
        flags=re.IGNORECASE,
    )
    if sentence_match:
        cleaned = sentence_match.group(1).strip()

    cleaned = re.split(
        r"\s+[-\u2013\u2014]\s+else\b|\s+else\s+delete\b",
        cleaned,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0].strip()
    return cleaned


def _looks_like_unfilled_optional_instruction(text):
    cleaned = (text or "").strip()
    if not cleaned:
        return True
    return bool(
        re.search(r"\bOPTIONAL\b|\bIF THE ROLE\b|\bname it\b|\belse delete\b", cleaned, re.IGNORECASE)
        or re.search(r"\bX\b.*\bY\b", cleaned)
        or "[" in cleaned
        or "]" in cleaned
    )


def _ensure_sentence_punctuation(text):
    cleaned = text.strip()
    if cleaned and cleaned[-1] not in ".!?":
        return f"{cleaned}."
    return cleaned


def _parse_cover_letter_rewrites(response):
    rewrites = {}
    for line in response.splitlines():
        m = re.match(r"^(\d+)\s*(?:[.)]|[-:])\s*(.+)$", line.strip())
        if m:
            idx = int(m.group(1)) - 1
            if idx >= 0:
                rewrites[idx] = m.group(2).strip()
    return rewrites


def _cover_letter_missing_rewrites_prompt(
    missing_items,
    *,
    company,
    title,
    intro,
    responsibilities,
    qualifications,
    company_reference=None,
    posting_context="direct_employer",
):
    company = _normalize_generated_text(company)
    company_reference = _normalize_generated_text(company_reference or company)
    context_rules = _cover_letter_agency_context_rules(posting_context)
    numbered_lines = "\n".join(
        f"{idx + 1}. {_normalize_generated_text(sentence)}"
        for idx, sentence in missing_items
    )
    return f"""Missing required cover-letter rewrites.

The previous response omitted these numbered sentence(s). Fill every missing sentence now.

Role: {title}
Company/folder name: {company}
Cover-letter employer reference (USE EXACTLY THIS): {company_reference}

Job description: {intro}
Responsibilities: {responsibilities}
Qualifications: {qualifications}

Sentences to fill:
{numbered_lines}
{context_rules}

Rules:
- Return each listed sentence IN FULL with ONLY the blank(s) replaced
- Use EXACTLY "{company_reference}" for company, employer, or client references; never modify it
- Use EXACTLY "{title}" for the role; never modify it
- Replace _ with employer/role-specific wording
- Replace any non-optional bracketed instruction with final cover-letter wording; never leave brackets in output
- For OPTIONAL blanks, return either the final sentence or DELETE
- Never use Unicode dash characters or double hyphens
- Return every listed number, preserving the original number
- Return ONLY the numbered sentences, nothing else"""


def _cover_letter_agency_context_rules(posting_context):
    if str(posting_context or "").strip().lower() != "agency_for_unknown_client":
        return ""
    return """

Agency/client context:
- This is an agency posting for an unnamed client.
- Use the cover-letter employer reference, not the posting company, when filling company/employer blanks.
- Do not describe the posting company as the employer.
- Avoid company/sector-specific claims unless the client context is stated in the job posting."""


def fill_cover_letter(
    path,
    company,
    title,
    intro,
    responsibilities,
    qualifications,
    *,
    company_reference=None,
    posting_context="direct_employer",
):
    doc = Document(path)
    company = _normalize_generated_text(company)
    company_reference = _normalize_generated_text(company_reference or company)
    title = _normalize_title_text(title)
    intro = _normalize_generated_text(intro)
    responsibilities = _normalize_generated_text(responsibilities)
    qualifications = _normalize_generated_text(qualifications)
    context_rules = _cover_letter_agency_context_rules(posting_context)
    now = datetime.now()
    today = f"{now.day} {now.strftime('%B %Y')}"

    # Replace dates
    date_replaced = False
    for para in _iter_all_paragraphs(doc):
        date_match = DATE_PATTERN.search(para.text)
        if date_match:
            _replace_para_text_range(para, date_match.start(), date_match.end(), today)
            print(f"  Date:          {today}")
            date_replaced = True
    if not date_replaced:
        print(f"  WARNING: no date found in template (today={today})")

    # Collect paragraphs with _ or [DESCRIPTION] blanks
    blank_paras = [
        (i, para, any(r.bold and _has_cover_letter_blank(r.text) for r in para.runs))
        for i, para in enumerate(_iter_all_paragraphs(doc))
        if _has_cover_letter_blank(para.text)
    ]

    if not blank_paras:
        doc.save(path)
        print("  Cover letter saved (no blanks found)")
        return

    # Extract only the sentence(s) containing _ or [DESCRIPTION] from each paragraph
    blank_items = []
    for _, para, had_bold_blank in blank_paras:
        sentences = _split_cover_letter_sentences(para.text)
        for si, s in enumerate(sentences):
            if _has_cover_letter_blank(s):
                blank_items.append((para, sentences, si, had_bold_blank))

    print(f"\n  Blanks found: {len(blank_items)}")

    numbered_lines = "\n".join(
        f"{j+1}. {_normalize_generated_text(sentences[si])}"
        for j, (_, sentences, si, _) in enumerate(blank_items)
    )

    prompt = f"""Fill in the blank(s) in each numbered sentence below. There are four blank types:
- _ (underscore): replace with company/role-specific content (e.g. company name, what the company does, why the applicant is drawn to this role)
- [DESCRIPTION]: replace the entire [DESCRIPTION] bracket (including the brackets) with content matching that description, written as a natural continuation of the sentence
- [INSTRUCTION TEXT]: Any other bracketed instruction is a required fill. Replace the entire bracketed instruction with final cover-letter wording that follows the instruction. Use examples inside the bracket only as hints; do not copy all examples.
- [OPTIONAL: ...] or OPTIONAL: ...: if the optional instruction is directly relevant to the role, replace the entire optional instruction with one concise natural sentence. If it is not directly relevant, return DELETE for that numbered item.

Role: {title}
Company/folder name: {company}
Cover-letter employer reference (USE EXACTLY THIS): {company_reference}

Job description: {intro}
Responsibilities: {responsibilities}
Qualifications: {qualifications}

Sentences to fill:
{numbered_lines}
{context_rules}

Rules:
- Return each sentence IN FULL with ONLY the blank(s) replaced; do NOT change, remove, or rephrase any other word
- Use EXACTLY "{company_reference}" for company, employer, or client references; never modify it
- Use EXACTLY "{title}" for the role; never modify it
- For [DESCRIPTION] blanks: replace the entire [DESCRIPTION] including brackets; never include brackets in output
- For [INSTRUCTION TEXT] blanks: replace the entire bracketed instruction including brackets; never include brackets in output
- Non-optional bracketed instructions are required fills; never return DELETE for them
- For OPTIONAL blanks: if the role mentions one of the optional examples, such as SAP, Oracle, IFRS, consolidations, or similar systems/standards, treat it as relevant and return one concrete final sentence
- For OPTIONAL blanks: never return the OPTIONAL instruction itself; return either the final sentence or DELETE
- Never include template labels or placeholders such as OPTIONAL, DESCRIPTION, [X], or [Y] in output
- Draw on the job description above to write specific, natural-sounding content; avoid generic filler
- Keep fills concise: one clause or short phrase for _, one sentence maximum for [DESCRIPTION] and [INSTRUCTION TEXT] blanks
- Never use Unicode dash characters or double hyphens. Use commas, parentheses, semicolons, or ordinary hyphens only inside compound words.
- Return every numbered sentence. Do not omit any number.
- Return ONLY the numbered sentences, nothing else"""

    response = call_llm(prompt, max_tokens=1000)
    rewrites = _parse_cover_letter_rewrites(response)
    missing_items = [
        (j, sentences[si])
        for j, (_, sentences, si, _) in enumerate(blank_items)
        if j not in rewrites
    ]
    if missing_items:
        repair_prompt = _cover_letter_missing_rewrites_prompt(
            missing_items,
            company=company,
            company_reference=company_reference,
            posting_context=posting_context,
            title=title,
            intro=intro,
            responsibilities=responsibilities,
            qualifications=qualifications,
        )
        rewrites.update(_parse_cover_letter_rewrites(call_llm(repair_prompt, max_tokens=800)))

    print(f"\n  Filled:")
    for j, (para, sentences, si, had_bold_blank) in enumerate(blank_items):
        raw_sentence = rewrites.get(j, "")
        if not raw_sentence:
            print(f"    {j+1}. WARNING: no rewrite returned")
            continue
        new_sentence, should_delete = _prepare_cover_letter_rewrite(
            raw_sentence,
            sentences[si],
        )
        if should_delete:
            print(f"    {j+1}. (deleted optional instruction)")
            _delete_text_once_preserving_runs(para, sentences[si])
            continue
        if not new_sentence:
            print(f"    {j+1}. WARNING: rewrite was empty after cleanup")
            continue
        print(f"    {j+1}. {new_sentence}")
        if not _replace_text_once_preserving_runs(para, sentences[si], new_sentence):
            print(f"    {j+1}. WARNING: original sentence not found in document")

    print()
    doc.save(path)
    print("  Cover letter saved")


# =============================================================================
# Resume Marker Tailoring
# =============================================================================

def fill_resume_markers(
    path,
    data,
    resume_source=None,
    category_dir=None,
    llm_call=call_llm,
    return_edit_records=False,
):
    """Tailor resume sentences wrapped in square brackets in a copied .docx."""
    doc = Document(path)
    marked_items = []
    for para in _iter_all_paragraphs(doc):
        sentence = _resume_editable_sentence(para.text)
        if sentence:
            marked_items.append((para, sentence))

    if not marked_items:
        return (0, []) if return_edit_records else 0

    numbered_sentences = "\n".join(
        f"{idx}. {sentence}"
        for idx, (_, sentence) in enumerate(marked_items, 1)
    )

    title = _normalize_title_text(data.get("title", ""))
    company = _normalize_generated_text(data.get("company", ""))
    intro = _normalize_generated_text(data.get("intro", ""))
    responsibilities = _normalize_generated_text(data.get("responsibilities", ""))
    qualifications = _normalize_generated_text(data.get("qualifications", ""))
    aggression = get_resume_tailoring_aggression()
    source_text = load_resume_source(resume_source, category_dir=category_dir)
    raw_extended_text = load_resume_extended(category_dir=category_dir)
    context_query = build_resume_context_query(
        data,
        [sentence for _, sentence in marked_items],
    )
    extended_selection = select_resume_extended_context(
        raw_extended_text,
        context_query,
        enabled=getattr(config, "RESUME_EXTENDED_SELECTION_ENABLED", True),
        max_sections=getattr(config, "RESUME_EXTENDED_MAX_SECTIONS", 8),
        max_chars=getattr(config, "RESUME_EXTENDED_MAX_CHARS", 12000),
        min_score=getattr(config, "RESUME_EXTENDED_MIN_SCORE", 2),
    )
    extended_text = extended_selection.text
    if category_dir:
        context_paths = get_resume_context_paths(category_dir=category_dir)
        for line in _resume_context_status(context_paths, extended_selection):
            print(line)
    source_block = (
        "Resume factual source. These are facts that may be directly claimed:\n"
        f"{source_text}\n\n"
        if source_text else ""
    )
    extended_block = (
        "Resume extended source. These are user-approved transferable mappings. "
        "Use them only as defensible phrasing guidance, not as new facts:\n"
        f"{extended_text}\n\n"
        if extended_text else ""
    )
    prompt = f"""Minimally tailor the bracketed resume sentences for this specific job.

Role: {title}
Company: {company}
Aggression mode: {aggression}
Aggression guidance: {_resume_aggression_guidance(aggression)}

Job description: {intro}
Responsibilities: {responsibilities}
Qualifications: {qualifications}

{source_block}{extended_block}Only use facts from the resume factual source, the current resume sentence, or the job description. Use the extended source only for user-approved transferable phrasing. If a useful keyword is not truthfully supported, do not add it.

Marked bracketed resume sentences:
{numbered_sentences}

Rules:
- Step 1: extract target ATS phrases from the job description and group them as high-priority or medium-priority
- Step 2: match each resume sentence to the strongest coherent keyword family before editing
- Return each numbered sentence IN FULL with the outer brackets removed
- Return EDIT or SKIP for each numbered sentence
- Use "N. EDIT | priority=<high|medium|low> | keyword=<keyword phrase> | sentence=<full revised sentence>" when a supported keyword fits coherently
- Use "N. SKIP | reason=<short reason>" when no coherent job-keyword fit exists
- Set priority=high for critical ATS/job matches, medium for useful matches, and low for weak/nice-to-have matches
- Make one compact role-specific keyword edit to every sentence when a supported job keyword can fit coherently
- Prefer replacing generic wording with exact job-description terminology over adding extra words
- Do not add a keyword just because it appears in the job description; it must fit the original sentence's evidence
- Do not count punctuation, slash replacement, or grammar-only cleanup as tailoring
- Keep the sentence structure and most original wording; add no more than one short phrase when substitution is not enough
- Prefer exact job-description finance terms and other domain-specific ATS keywords when they are supported by the factual source, current sentence, or extended source
- If no supported job-specific keyword can fit truthfully and coherently, return the original sentence without brackets
- Do not invent skills, employers, credentials, responsibilities, metrics, tools, or outcomes
- Preserve high-signal technical wording, metrics, tools, company names, acronyms, and senior finance/domain language
- Do not simplify, de-jargonize, or convert technical finance/domain language into generic plain English
- Do not use hidden text, keyword stuffing, tables, or formatting tricks
- Return ONLY numbered EDIT/SKIP lines, nothing else"""

    response = llm_call(prompt, max_tokens=2000)

    rewrites = {}
    rewrite_keywords = {}
    rewrite_priorities = {}
    skip_reasons = {}
    for line in response.splitlines():
        parsed = _parse_resume_response_line(line)
        if not parsed:
            continue
        idx, action, keyword, body, priority = parsed
        if action == "SKIP":
            skip_reasons[idx] = body or "no coherent job-keyword fit"
        else:
            rewrites[idx] = _clean_resume_rewrite(body)
            rewrite_keywords[idx] = keyword
            rewrite_priorities[idx] = priority

    changed = 0
    edits = []
    edit_records = []
    skipped = []
    for idx, (para, sentence) in enumerate(marked_items, 1):
        if idx in skip_reasons:
            new_text = sentence
            skipped.append((idx, skip_reasons[idx]))
        else:
            new_text = rewrites.get(idx) or sentence
            if idx not in rewrites:
                skipped.append((idx, "no rewrite returned"))
        if new_text != sentence:
            edits.append((idx, rewrite_keywords.get(idx, ""), new_text))
            edit_records.append(
                {
                    "marker": idx,
                    "keyword": rewrite_keywords.get(idx, ""),
                    "priority": rewrite_priorities.get(idx, "medium"),
                    "original": sentence,
                    "edited": new_text,
                }
            )
        _set_para_text(para, new_text)
        changed += 1

    _print_resume_summary(
        markers=len(marked_items),
        aggression=aggression,
        edits=edits,
        skipped=skipped,
        verbosity=get_cli_verbosity(),
    )

    doc.save(path)
    print("  Resume saved")
    if return_edit_records:
        return changed, edit_records
    return changed


def _parse_resume_response_line(line):
    cleaned = line.strip()
    m = re.match(r"^(\d+)\.\s*(.+)$", cleaned)
    if not m:
        return None

    idx = int(m.group(1))
    body = m.group(2).strip()

    pipe = re.match(r"^(EDIT|SKIP)\s*\|\s*(.+)$", body, re.IGNORECASE)
    if pipe:
        action = pipe.group(1).upper()
        fields = _parse_pipe_fields(pipe.group(2))
        if action == "SKIP":
            return idx, action, "", fields.get("reason", ""), ""
        return (
            idx,
            action,
            fields.get("keyword", ""),
            fields.get("sentence", ""),
            _normalize_resume_priority(fields.get("priority", "medium")),
        )

    legacy = re.match(r"^(EDIT|SKIP)\s*:\s*(.+)$", body, re.IGNORECASE)
    if legacy:
        action = legacy.group(1).upper()
        return idx, action, "", legacy.group(2).strip(), "medium"

    return idx, "EDIT", "", body, "medium"


def _parse_pipe_fields(text):
    fields = {}
    for part in re.split(r"\s+\|\s+", text.strip()):
        if "=" not in part:
            continue
        key, value = part.split("=", 1)
        fields[key.strip().lower()] = value.strip()
    return fields


def _normalize_resume_priority(priority):
    priority = str(priority or "medium").strip().lower()
    if priority not in {"high", "medium", "low"}:
        return "medium"
    return priority


def _print_resume_summary(markers, aggression, edits, skipped, verbosity):
    skipped_count = markers - len(edits)
    if verbosity == "quiet":
        print(f"\n  Resume: markers={markers} edited={len(edits)} skipped={skipped_count}")
        return

    print("\n  Resume")
    print(f"    Markers: {markers} | Edited: {len(edits)} | Skipped: {skipped_count}")
    print(f"    Aggression: {aggression}")

    print("\n    Edits:")
    if edits:
        for idx, keyword, sentence in edits:
            label = f"+ {keyword}" if keyword else "edited"
            print(f"      {idx}. {label}")
            if verbosity == "debug":
                print(f"         {sentence}")
    else:
        print("      none")

    if verbosity == "debug" and skipped:
        print("\n    Skipped:")
        for idx, reason in skipped:
            print(f"      {idx}. {reason}")


def _resume_editable_sentence(text):
    """Return the sentence inside an editable outer bracket marker."""
    cleaned = text.strip()
    match = RESUME_MARKER_PATTERN.match(cleaned)
    if not match:
        return None
    sentence = match.group(1).strip()
    if RESUME_CONTROL_LABEL_PATTERN.match(sentence):
        return None
    if "[" in sentence or "]" in sentence:
        return None
    return sentence


def _clean_resume_rewrite(text):
    """Remove accidental outer bracket markers from a model rewrite."""
    cleaned = text.strip()
    sentence = _resume_editable_sentence(cleaned)
    if sentence:
        cleaned = sentence
    cleaned = re.sub(r"\s+([.,;:])", r"\1", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return _normalize_generated_text(cleaned)


# =============================================================================
# Page Fit And PDF Conversion
# =============================================================================

def fit_docx_to_page_limit(
    docx_path,
    pdf_path,
    label,
    limit,
    *,
    llm_call=call_llm,
    convert_fn=None,
    page_count_fn=None,
    max_attempts=None,
):
    """Render a DOCX, shorten if it exceeds the page limit, and retry."""
    if convert_fn is None:
        convert_fn = convert_docx_to_pdf
    if page_count_fn is None:
        page_count_fn = _pdf_page_count
    if max_attempts is None:
        max_attempts = get_page_fit_max_attempts()

    shortened_indices = set()
    attempts = max_attempts + 1
    for attempt in range(1, attempts + 1):
        attempt_pdf = _page_fit_attempt_pdf_path(pdf_path, attempt)
        convert_fn(docx_path, attempt_pdf)
        pages = page_count_fn(attempt_pdf)

        if pages <= limit:
            _replace_pdf(attempt_pdf, pdf_path)
            print(f"  {label} PDF: {pdf_path}")
            if attempt > 1:
                print(f"  Page fit: {label} shortened to {pages} page(s)")
            return True

        if attempt > max_attempts:
            _replace_pdf(attempt_pdf, pdf_path)
            print(f"  {label} PDF: {pdf_path}")
            print(f"  WARNING: {label} is {pages} pages (limit: {limit}) after {max_attempts} shortening attempt(s)")
            return False

        print(f"  Page fit: {label} is {pages} pages (limit: {limit}); shortening attempt {attempt}/{max_attempts}")
        changed, changed_indices = shorten_docx_for_page_limit(
            docx_path,
            label,
            pages,
            limit,
            llm_call=llm_call,
            exclude_indices=shortened_indices,
            return_changed_indices=True,
        )
        if not changed:
            _replace_pdf(attempt_pdf, pdf_path)
            print(f"  {label} PDF: {pdf_path}")
            print(f"  WARNING: {label} is {pages} pages (limit: {limit}); no safe shortening changes were accepted")
            return False
        shortened_indices.update(changed_indices)
        _discard_file(attempt_pdf)

    return False


def fit_resume_docx_to_page_limit(
    docx_path,
    pdf_path,
    limit,
    edit_records,
    *,
    convert_fn=None,
    page_count_fn=None,
):
    """Render a tailored resume, reverting lowest-priority edits until it fits."""
    if convert_fn is None:
        convert_fn = convert_docx_to_pdf
    if page_count_fn is None:
        page_count_fn = _pdf_page_count

    attempt = 1
    attempt_pdf = _page_fit_attempt_pdf_path(pdf_path, attempt)
    convert_fn(docx_path, attempt_pdf)
    pages = page_count_fn(attempt_pdf)
    if pages <= limit:
        _replace_pdf(attempt_pdf, pdf_path)
        print(f"  Resume PDF: {pdf_path}")
        return True

    print(f"  Page fit: Resume is {pages} pages (limit: {limit}); reverting lowest-priority edits")
    reverted = []
    for record in _resume_rollback_order(edit_records):
        if not _revert_resume_edit(docx_path, record):
            continue
        reverted.append(record)
        keyword = record.get("keyword") or "edited keyword"
        print(f"    Reverted {record.get('marker')}. {keyword}")

        _discard_file(attempt_pdf)
        attempt += 1
        attempt_pdf = _page_fit_attempt_pdf_path(pdf_path, attempt)
        convert_fn(docx_path, attempt_pdf)
        pages = page_count_fn(attempt_pdf)
        if pages <= limit:
            _replace_pdf(attempt_pdf, pdf_path)
            print(f"  Resume PDF: {pdf_path}")
            print(f"  Page fit: Resume fit after reverting {len(reverted)} edit(s)")
            return True

    _replace_pdf(attempt_pdf, pdf_path)
    print(f"  Resume PDF: {pdf_path}")
    print(f"  WARNING: Resume is still {pages} pages (limit: {limit}) after reverting {len(reverted)} edit(s)")
    return False


def _resume_rollback_order(edit_records):
    priority_rank = {"low": 0, "medium": 1, "high": 2}

    def key(record):
        priority = _normalize_resume_priority(record.get("priority", "medium"))
        original = record.get("original", "")
        edited = record.get("edited", "")
        added_chars = len(edited) - len(original)
        return (priority_rank[priority], -added_chars, record.get("marker", 0))

    return sorted(edit_records or [], key=key)


def _revert_resume_edit(docx_path, record):
    edited = record.get("edited", "")
    original = record.get("original", "")
    if not edited or not original or edited == original:
        return False

    doc = Document(docx_path)
    for para in _iter_all_paragraphs(doc):
        if para.text.strip() == edited:
            _set_para_text(para, original)
            doc.save(docx_path)
            return True
    return False


def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        import win32com.client
    except ImportError:
        convert(docx_path, pdf_path)
        return

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(
            str(docx_path),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        try:
            doc.ExportAsFixedFormat(str(pdf_path), 17)
        finally:
            doc.Close(False)
    finally:
        word.Quit()


def _pdf_page_count(pdf_path):
    return len(PdfReader(pdf_path).pages)


def _page_fit_attempt_pdf_path(pdf_path, attempt):
    root, ext = os.path.splitext(pdf_path)
    return f"{root}.pagefit_attempt_{attempt}{ext}"


def _replace_pdf(src, dst):
    if os.path.abspath(src) == os.path.abspath(dst):
        return
    _discard_file(dst)
    os.replace(src, dst)


def _discard_file(path):
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except OSError:
        pass


# =============================================================================
# Cover Letter Page-Fit Shortening Helpers
# =============================================================================

def shorten_docx_for_page_limit(
    docx_path,
    label,
    pages,
    limit,
    *,
    llm_call=call_llm,
    exclude_indices=None,
    return_changed_indices=False,
):
    doc = Document(docx_path)
    min_ratio = get_page_fit_min_line_retain_ratio()
    items = _shortenable_paragraph_items(
        doc,
        exclude_indices=exclude_indices,
        max_lines=get_page_fit_max_lines_per_attempt(),
    )
    if not items:
        return (0, []) if return_changed_indices else 0

    numbered_lines = "\n".join(
        f"{idx}. {_normalize_generated_text(para.text.strip())}"
        for idx, para in items
    )
    retain_percent = int(min_ratio * 100)
    prompt = f"""Micro-shorten these {label} lines so the document can fit within {limit} page(s).

Current rendered PDF page count: {pages}
Target page count: {limit}
Maximum candidate lines this attempt: {len(items)}
Each replacement must retain at least {retain_percent}% of the original line length.

Lines:
{numbered_lines}

Rules:
- Return slight replacements for the numbered lines only
- Shorten only enough to save space; do not rewrite the whole line
- Preserve facts, metrics, tools, company names, credentials, dates, and role titles
- Remove filler, repeated wording, and weak modifiers before removing useful keywords
- Keep ATS/job-specific keywords that are already present
- Do not add new facts or new claims
- Prefer small wording compression over changing meaning
- It is acceptable to return the original line when a safe slight shortening is not possible
- Return ONLY numbered lines, nothing else"""

    response = llm_call(prompt, max_tokens=1200)
    rewrites = {}
    for line in response.splitlines():
        m = re.match(r"^(\d+)\.\s*(.+)$", line.strip())
        if m:
            rewrites[int(m.group(1))] = _normalize_generated_text(m.group(2).strip())

    changed = 0
    changed_indices = []
    rejected = []
    for idx, para in items:
        old_text = para.text.strip()
        new_text = rewrites.get(idx)
        if new_text and _is_acceptable_shortening(old_text, new_text, min_ratio):
            _set_para_text(para, new_text)
            changed += 1
            changed_indices.append(idx)
        elif new_text and new_text != old_text:
            rejected.append((idx, old_text, new_text))

    if not changed and rejected:
        repair_response = llm_call(
            _page_fit_repair_prompt(label, limit, rejected, min_ratio),
            max_tokens=1200,
        )
        repair_rewrites = {}
        for line in repair_response.splitlines():
            m = re.match(r"^(\d+)\.\s*(.+)$", line.strip())
            if m:
                repair_rewrites[int(m.group(1))] = _normalize_generated_text(m.group(2).strip())

        for idx, para in items:
            old_text = para.text.strip()
            new_text = repair_rewrites.get(idx)
            if new_text and _is_acceptable_shortening(old_text, new_text, min_ratio):
                _set_para_text(para, new_text)
                changed += 1
                changed_indices.append(idx)

    if changed:
        doc.save(docx_path)
    if return_changed_indices:
        return changed, changed_indices
    return changed


def _shortenable_paragraph_items(doc, exclude_indices=None, max_lines=None):
    exclude_indices = set(exclude_indices or [])
    items = []
    for idx, para in enumerate(_iter_all_paragraphs(doc), 1):
        text = para.text.strip()
        if idx not in exclude_indices and _is_shortenable_paragraph(text):
            items.append((idx, para))
    items = sorted(items, key=lambda item: len(item[1].text.strip()), reverse=True)
    if max_lines:
        items = items[:max_lines]
    items = sorted(items, key=lambda item: item[0])
    return items


def _is_acceptable_shortening(original, candidate, min_retain_ratio):
    candidate = candidate.strip()
    if not candidate or candidate == original:
        return False
    if len(candidate) >= len(original):
        return False
    return len(candidate) >= int(len(original) * min_retain_ratio)


def _page_fit_repair_prompt(label, limit, rejected, min_ratio):
    lines = []
    retain_percent = int(min_ratio * 100)
    for idx, original, rejected_text in rejected:
        minimum = int(len(original) * min_ratio)
        maximum = len(original) - 1
        lines.append(
            f"{idx}. target={minimum}-{maximum} characters | original={original}"
        )
    numbered_lines = "\n".join(lines)
    return f"""Previous replacements were rejected because they cut too much text.

Rewrite these {label} lines again so the document can fit within {limit} page(s).

Each replacement must be shorter than the original but retain at least {retain_percent}% of the original line length.
Use the target character range shown for each line.

Lines:
{numbered_lines}

Rules:
- Return replacements for the numbered lines only
- Keep facts, metrics, tools, company names, credentials, dates, role titles, and ATS keywords
- Make only micro-edits: remove filler, repeated wording, or weak modifiers
- Do not simplify technical wording into generic language
- Do not add new facts or claims
- Return ONLY numbered lines, nothing else"""


def _is_shortenable_paragraph(text):
    if len(text) < 70:
        return False
    if text.isupper():
        return False
    lowered = text.lower()
    if "email:" in lowered or "mobile:" in lowered or "linkedin:" in lowered:
        return False
    return True


# =============================================================================
# Bundle Generation
# =============================================================================

def _merge_cover_letter_bundle(cover_pdf: str, output_folder: str):
    """Merge cover letter + BUNDLE_APPENDIX into one PDF in output_folder."""
    bundle_path = os.path.join(output_folder, f"{BUNDLE_NAME}.pdf")
    writer = PdfWriter()
    for path in [cover_pdf] + BUNDLE_APPENDIX:
        if not os.path.exists(path):
            print(f"  WARNING: bundle file not found, skipping: {path}")
            continue
        writer.append(path)
    with open(bundle_path, "wb") as f:
        writer.write(f)
    print(f"  Bundle PDF: {bundle_path}")


# =============================================================================
# Application Orchestration
# =============================================================================

def _cover_letter_company_reference(data):
    context = str(data.get("posting_context") or "direct_employer").strip().lower()
    reference = str(data.get("cover_letter_company_reference") or "").strip()
    if context == "agency_for_unknown_client":
        return reference or "your client"
    return reference or data.get("company", "")


def generate_application(data, category=None):
    title = clean_job_title(data["title"])
    company = data["company"]
    company_reference = _cover_letter_company_reference(data)
    posting_context = data.get("posting_context", "direct_employer")
    generation_warnings = []

    if category is None:
        category = classify_job(title, data["description"])
    resume_pdf, resume_docx, cover_docx, category_dir = get_paths(category)

    folder_name = re.sub(r'[<>:"/\\|?*]', '-', f"{company} - {title}")
    output_folder = os.path.join(config.OUTPUT_BASE, folder_name)
    os.makedirs(output_folder, exist_ok=True)

    resume_pdf_dest = os.path.join(output_folder, os.path.basename(resume_pdf))
    shutil.copy(resume_pdf, resume_pdf_dest)

    if resume_docx:
        resume_docx_dest = os.path.join(output_folder, os.path.basename(resume_docx))
        shutil.copy(resume_docx, resume_docx_dest)
        resume_markers, resume_edit_records = fill_resume_markers(
            resume_docx_dest,
            data,
            category_dir=category_dir,
            return_edit_records=True,
        )
        if resume_markers:
            tailored_resume_pdf = get_tailored_resume_pdf_path(resume_docx_dest)
            try:
                fit_resume_docx_to_page_limit(
                    resume_docx_dest,
                    tailored_resume_pdf,
                    get_resume_page_limit(),
                    resume_edit_records,
                )
            except Exception as e:
                print(f"  WARNING: resume PDF conversion failed ({e})")
                print("  Make sure Microsoft Word is installed and the .docx is not open.")
                generation_warnings.append(f"Resume PDF conversion failed: {e}")

    cover_dest = os.path.join(output_folder, os.path.basename(cover_docx))
    shutil.copy(cover_docx, cover_dest)

    # Write position description PDF (captured during scraping while session was live)
    position_pdf = os.path.join(output_folder, "Position Description.pdf")
    if data.get("pdf_bytes"):
        with open(position_pdf, "wb") as f:
            f.write(data["pdf_bytes"])
        print(f"  Position description PDF: {position_pdf}")

    fill_cover_letter(
        cover_dest, company, title,
        data.get("intro", ""),
        data.get("responsibilities", ""),
        data.get("qualifications", ""),
        company_reference=company_reference,
        posting_context=posting_context,
    )

    pdf_dest = get_cover_letter_pdf_path(cover_dest)
    try:
        fit_docx_to_page_limit(
            cover_dest,
            pdf_dest,
            "Cover letter",
            get_cover_letter_page_limit(),
        )
        if BUNDLE_APPENDIX:
            _merge_cover_letter_bundle(pdf_dest, output_folder)
    except Exception as e:
        print(f"  WARNING: PDF conversion failed ({e})")
        print("  Make sure Microsoft Word is installed and the .docx is not open.")
        generation_warnings.append(f"Cover letter PDF conversion failed: {e}")

    if generation_warnings:
        print(f"\nDone with warnings! Saved to: {output_folder}")
        print("  Issues:")
        for warning in generation_warnings:
            print(f"    - {warning}")
    else:
        print(f"\nDone! Saved to: {output_folder}")
    os.startfile(output_folder)
    return output_folder
